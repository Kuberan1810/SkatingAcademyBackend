import csv
import io
import os
import re
import shutil

from datetime import date, datetime
from pathlib import Path
from typing import Any

import cv2
import numpy as np
from docx import Document
from openpyxl import load_workbook

from PIL import (
    Image,
    ImageEnhance,
    ImageOps,
    UnidentifiedImageError,
)

import pytesseract

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

from sqlalchemy.orm import Session

from app.models.student import Student


# =========================================================
# PADDLE OCR (RAPIDOCR ONNX) CONFIGURATION
# Lightweight, pure-Python / pip-installable, no Docker required
# =========================================================

_rapid_ocr_engine = None


def get_ocr_engine():
    """
    Returns the singleton RapidOCR (PaddleOCR ONNX) engine instance.
    """
    global _rapid_ocr_engine
    if _rapid_ocr_engine is None:
        try:
            from rapidocr_onnxruntime import RapidOCR
            _rapid_ocr_engine = RapidOCR()
        except ImportError:
            _rapid_ocr_engine = None
    return _rapid_ocr_engine


def get_tesseract_cmd() -> str:
    return "rapidocr-onnxruntime (PaddleOCR)"


def configure_tesseract() -> str:
    return "rapidocr-onnxruntime (PaddleOCR)"

# =========================================================
# FIELD ALIASES
# =========================================================

FIELD_ALIASES = {

    # -----------------------------------------------------
    # REQUIRED
    # -----------------------------------------------------

    "full name": "full_name",
    "fullname": "full_name",
    "full_name": "full_name",
    "name": "full_name",
    "student name": "full_name",
    "student_name": "full_name",

    "dob": "dob",
    "date of birth": "dob",
    "dateofbirth": "dob",
    "date_of_birth": "dob",

    "gender": "gender",
    "sex": "gender",

    "parent name": "parent_name",
    "parentname": "parent_name",
    "parent_name": "parent_name",
    "guardian name": "parent_name",
    "guardian": "parent_name",

    "phone number": "phone_number",
    "phonenumber": "phone_number",
    "phone_number": "phone_number",
    "phone": "phone_number",
    "mobile": "phone_number",
    "mobile number": "phone_number",

    "monthly fee": "monthly_fee",
    "monthlyfee": "monthly_fee",
    "monthly_fee": "monthly_fee",
    "fee": "monthly_fee",

    # -----------------------------------------------------
    # OPTIONAL
    # -----------------------------------------------------

    "blood group": "blood_group",
    "bloodgroup": "blood_group",
    "blood_group": "blood_group",

    "emergency contact": "emergency_contact",
    "emergencycontact": "emergency_contact",
    "emergency_contact": "emergency_contact",
    "emergency phone": "emergency_contact",

    "join date": "join_date",
    "joindate": "join_date",
    "join_date": "join_date",

    "avatar uri": "avatar_uri",
    "avatar_uri": "avatar_uri",
    "avatar": "avatar_uri",
}


# =========================================================
# TEXT HELPERS
# =========================================================

def clean_text(
    value: Any,
) -> str:

    if value is None:
        return ""

    return str(
        value
    ).strip()


def normalize_key(
    value: Any,
) -> str:

    value = clean_text(
        value
    ).lower()

    value = value.replace(
        "_",
        " ",
    )

    value = value.replace(
        "-",
        " ",
    )

    value = re.sub(
        r"\s+",
        " ",
        value,
    )

    return value.strip()


def map_field_name(
    value: Any,
) -> str | None:

    return FIELD_ALIASES.get(
        normalize_key(
            value
        )
    )


# =========================================================
# DATE
# =========================================================

def parse_date_value(
    value: Any,
) -> date | None:

    if value is None:
        return None

    if isinstance(
        value,
        datetime,
    ):
        return value.date()

    if isinstance(
        value,
        date,
    ):
        return value

    value = clean_text(
        value
    )

    if not value:
        return None

    formats = [
        "%Y-%m-%d",
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%m/%d/%Y",
        "%d.%m.%Y",
        "%Y/%m/%d",
        "%d %b %Y",
        "%d %B %Y",
    ]

    for fmt in formats:
        try:
            return datetime.strptime(
                value,
                fmt,
            ).date()

        except ValueError:
            continue

    return None


# =========================================================
# INTEGER
# =========================================================

def parse_int_value(
    value: Any,
) -> int | None:

    if value is None:
        return None

    text = clean_text(
        value
    )

    if not text:
        return None

    text = re.sub(
        r"[₹,\s]",
        "",
        text,
    )

    try:
        return int(
            float(text)
        )

    except (
        TypeError,
        ValueError,
    ):
        return None


# =========================================================
# PHONE
# =========================================================

def clean_phone(
    value: Any,
) -> str:

    if value is None:
        return ""

    return re.sub(
        r"[^\d+]",
        "",
        str(value),
    )


# =========================================================
# NORMALIZE RECORD
# =========================================================

def normalize_record(
    record: dict[str, Any],
    db: Session,
) -> dict[str, Any]:

    errors: list[str] = []
    warnings: list[str] = []

    normalized: dict[str, Any] = {}

    # =====================================================
    # MAP INPUT HEADERS
    # =====================================================

    for raw_key, value in record.items():

        field_name = map_field_name(
            raw_key
        )

        if field_name:

            normalized[
                field_name
            ] = value

    # =====================================================
    # REQUIRED FIELDS (Only Name, DOB, Phone mandatory)
    # =====================================================

    full_name = clean_text(
        normalized.get(
            "full_name"
        )
    )

    dob = parse_date_value(
        normalized.get(
            "dob"
        )
    )

    phone_number = clean_phone(
        normalized.get(
            "phone_number"
        )
    )

    # =====================================================
    # OPTIONAL FIELDS
    # =====================================================

    gender = (
        clean_text(
            normalized.get(
                "gender"
            )
        ).title()
        or None
    )

    parent_name = (
        clean_text(
            normalized.get(
                "parent_name"
            )
        )
        or None
    )

    monthly_fee = parse_int_value(
        normalized.get(
            "monthly_fee"
        )
    )

    blood_group = (
        clean_text(
            normalized.get(
                "blood_group"
            )
        ).upper()
        or None
    )

    emergency_contact = (
        clean_phone(
            normalized.get(
                "emergency_contact"
            )
        )
        or None
    )

    join_date = parse_date_value(
        normalized.get(
            "join_date"
        )
    )

    avatar_uri = (
        clean_text(
            normalized.get(
                "avatar_uri"
            )
        )
        or None
    )

    # =====================================================
    # REQUIRED VALIDATION (Only Name, DOB, Phone required)
    # =====================================================

    if not full_name:

        errors.append(
            "Full name is required"
        )

    if dob is None:

        errors.append(
            "Valid date of birth is required"
        )

    if not phone_number:

        errors.append(
            "Phone number is required"
        )

    elif len(
        re.sub(
            r"\D",
            "",
            phone_number,
        )
    ) < 10:

        errors.append(
            "Phone number must contain at least 10 digits"
        )

    # =====================================================
    # OPTIONAL VALIDATION
    # =====================================================

    if gender and gender not in {
        "Male",
        "Female",
        "Other",
    }:

        errors.append(
            "Gender must be Male, Female, or Other"
        )

    if monthly_fee is not None and monthly_fee < 0:

        errors.append(
            "Monthly fee cannot be negative"
        )

    # =====================================================
    # OPTIONAL VALIDATION
    # =====================================================

    if blood_group:

        if blood_group not in {
            "A+",
            "A-",
            "B+",
            "B-",
            "AB+",
            "AB-",
            "O+",
            "O-",
        }:

            errors.append(
                "Invalid blood group"
            )

    if emergency_contact:

        emergency_digits = re.sub(
            r"\D",
            "",
            emergency_contact,
        )

        if len(
            emergency_digits
        ) < 10:

            errors.append(
                "Emergency contact must contain at least 10 digits"
            )

    # =====================================================
    # JOIN DATE
    # =====================================================
    #
    # Optional.
    # Keep null during preview.
    # Confirm API defaults to today.
    #
    # =====================================================

    # =====================================================
    # DUPLICATE CHECK
    # =====================================================

    if phone_number:

        duplicate = (
            db.query(
                Student
            )
            .filter(
                Student.phone_number
                == phone_number
            )
            .first()
        )

        if duplicate:

            warnings.append(
                (
                    "Student already exists "
                    "with this phone number"
                )
            )

    # =====================================================
    # STATUS
    # =====================================================

    if errors:

        status = "invalid"

    elif warnings:

        status = "warning"

    else:

        status = "valid"

    # =====================================================
    # RETURN
    # =====================================================

    return {

        "full_name":
            full_name,

        "dob":
            dob,

        "gender":
            gender,

        "blood_group":
            blood_group,

        "join_date":
            join_date,

        "parent_name":
            parent_name,

        "phone_number":
            phone_number,

        "emergency_contact":
            emergency_contact,

        "monthly_fee":
            monthly_fee,

        "avatar_uri":
            avatar_uri,

        "status":
            status,

        "errors":
            errors,

        "warnings":
            warnings,
    }


# =========================================================
# XLSX
# =========================================================

def parse_xlsx(
    file_path: str,
):

    workbook = load_workbook(
        file_path,
        data_only=True,
    )

    sheet = workbook.active

    rows = list(
        sheet.iter_rows(
            values_only=True
        )
    )

    if not rows:
        return []

    headers = [
        clean_text(
            value
        )
        for value
        in rows[0]
    ]

    records = []

    for row in rows[1:]:

        if not any(
            value not in (
                None,
                "",
            )
            for value
            in row
        ):
            continue

        record = {}

        for index, value in enumerate(
            row
        ):

            if (
                index
                < len(headers)
                and headers[index]
            ):

                record[
                    headers[index]
                ] = value

        records.append(
            record
        )

    return records


# =========================================================
# CSV
# =========================================================

def parse_csv(
    file_path: str,
):

    with open(
        file_path,
        "r",
        encoding="utf-8-sig",
        newline="",
    ) as file:

        reader = csv.DictReader(
            file
        )

        return [
            dict(row)
            for row in reader
            if any(
                clean_text(value)
                for value
                in row.values()
            )
        ]


# =========================================================
# DOCX
# =========================================================

def parse_docx(
    file_path: str,
):

    document = Document(
        file_path
    )

    # -----------------------------------------------------
    # TABLE
    # -----------------------------------------------------

    for table in document.tables:

        if not table.rows:
            continue

        headers = [
            clean_text(
                cell.text
            )
            for cell
            in table.rows[0].cells
        ]

        records = []

        for row in table.rows[1:]:

            values = [
                clean_text(
                    cell.text
                )
                for cell
                in row.cells
            ]

            if not any(values):
                continue

            record = {}

            for index, value in enumerate(
                values
            ):

                if (
                    index
                    < len(headers)
                    and headers[index]
                ):

                    record[
                        headers[index]
                    ] = value

            records.append(
                record
            )

        if records:
            return records

    # -----------------------------------------------------
    # TEXT
    # -----------------------------------------------------

    text = "\n".join(
        paragraph.text
        for paragraph
        in document.paragraphs
        if paragraph.text.strip()
    )

    return parse_text_records(
        text
    )


# =========================================================
# TEXT PARSER
# =========================================================

def parse_text_records(
    text: str,
) -> list[dict[str, Any]]:
    text = (
        text
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )

    lines = [
        line.strip()
        for line
        in text.split("\n")
        if line.strip()
    ]

    if not lines:
        return []

    records = []
    current = {}
    pending_field = None

    field_pattern = re.compile(
        r"^\s*(.+?)\s*[:=]\s*(.*?)\s*$"
    )

    for line in lines:

        # -------------------------------------------------
        # Student 1 / Student 2 separator
        # -------------------------------------------------
        if re.match(
            r"^(?:student|record|candidate|entry)\s*\d+",
            line,
            re.IGNORECASE,
        ):
            if current:
                records.append(current)
                current = {}
            pending_field = None
            continue

        match = field_pattern.match(line)
        if match:
            raw_key = match.group(1).strip()
            value = match.group(2).strip()
            field_name = map_field_name(raw_key)

            if field_name:
                if value:
                    if field_name in current and field_name == "full_name":
                        records.append(current)
                        current = {}
                    current[field_name] = value
                    pending_field = None
                else:
                    # Key without value on this line (e.g. "Name:" then "Ramesh")
                    pending_field = field_name
            continue

        # Handle pending field from previous line
        if pending_field:
            if pending_field in current and pending_field == "full_name":
                records.append(current)
                current = {}
            current[pending_field] = line
            pending_field = None
            continue

    if current:
        records.append(current)

    return records


def parse_txt(
    file_path: str,
):

    with open(
        file_path,
        "r",
        encoding="utf-8-sig",
    ) as file:

        return parse_text_records(
            file.read()
        )


# =========================================================
# PADDLE OCR (RAPIDOCR ONNX) TEXT EXTRACTION
# =========================================================

def extract_text_from_image(
    image_input: str | bytes | Image.Image | np.ndarray,
) -> str:
    """
    Extracts text from an image using PaddleOCR (RapidOCR ONNX).
    - 100% pure Python / pip package (no apt-get or Docker needed)
    - Lightweight, fast (~50ms), and highly accurate on tables and mobile photos.
    """
    img_np = None

    try:
        if isinstance(image_input, np.ndarray):
            img_np = image_input
        elif isinstance(image_input, Image.Image):
            img_np = np.array(image_input.convert("RGB"))
        elif isinstance(image_input, (bytes, bytearray)):
            np_arr = np.frombuffer(image_input, np.uint8)
            img_np = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
            if img_np is None:
                pil_img = Image.open(io.BytesIO(image_input)).convert("RGB")
                img_np = np.array(pil_img)
        elif isinstance(image_input, str):
            if os.path.isfile(image_input):
                img_np = cv2.imread(image_input)
                if img_np is None:
                    pil_img = Image.open(image_input).convert("RGB")
                    img_np = np.array(pil_img)
    except Exception:
        img_np = None

    if img_np is None:
        raise ValueError("Unable to read or decode image file.")

    # 1. Try RapidOCR (PaddleOCR ONNX)
    engine = get_ocr_engine()
    if engine is not None:
        try:
            results, _ = engine(img_np)
            if results:
                # Sort text blocks top-to-bottom
                results_sorted = sorted(results, key=lambda item: item[0][0][1])
                lines = [item[1].strip() for item in results_sorted if item[1].strip()]
                extracted_text = "\n".join(lines)
                if extracted_text.strip():
                    return extracted_text
        except Exception:
            pass

    # 2. Optional Fallback to PyTesseract if available
    try:
        pil_fallback = Image.fromarray(img_np)
        extracted = pytesseract.image_to_string(pil_fallback, config="--psm 6")
        if extracted.strip():
            return extracted
    except Exception:
        pass

    return ""


def preprocess_image(image: Image.Image) -> Image.Image:
    return image


def parse_image(
    file_path_or_bytes: str | bytes | Image.Image,
) -> list[dict[str, Any]]:
    """
    Extracts structured student records from image files (.jpg, .jpeg, .png, .webp)
    using lightweight PaddleOCR.
    """
    try:
        text = extract_text_from_image(file_path_or_bytes)
    except UnidentifiedImageError as exc:
        raise ValueError("Invalid or unreadable image file format.") from exc
    except Exception as exc:
        raise ValueError(f"Failed to process image for OCR: {exc}") from exc

    if not text.strip():
        raise ValueError(
            "No readable student data found in image. "
            "Please ensure the image is clear and legible."
        )

    return parse_text_records(text)


# =========================================================
# PDF
# =========================================================

def parse_pdf(
    file_path: str,
):
    """
    Robust PDF student import.

    Flow:
    1. Try PyMuPDF text extraction.
    2. If text exists -> parse text.
    3. If PDF has no selectable text -> OCR pages with PaddleOCR.
    4. If both fail -> return a clear error.
    """

    # =====================================================
    # 1. PYMuPDF
    # =====================================================

    try:
        import fitz

        document = fitz.open(
            file_path
        )

        extracted_text_parts = []

        for page_number in range(
            document.page_count
        ):
            page = document.load_page(
                page_number
            )

            page_text = (
                page.get_text(
                    "text"
                )
                or ""
            )

            if page_text.strip():
                extracted_text_parts.append(
                    page_text
                )

        text = "\n".join(
            extracted_text_parts
        )

        # -------------------------------------------------
        # Normal text PDF
        # -------------------------------------------------

        if text.strip():

            document.close()

            records = parse_text_records(
                text
            )

            if records:
                return records

        # =================================================
        # 2. OCR FALLBACK (PaddleOCR)
        # =================================================

        ocr_records = []

        for page_number in range(
            document.page_count
        ):
            page = document.load_page(
                page_number
            )

            # Render PDF page to image at 2x resolution
            pix = page.get_pixmap(
                matrix=fitz.Matrix(
                    2,
                    2,
                ),
                alpha=False,
            )

            img_np = np.frombuffer(
                pix.samples,
                dtype=np.uint8,
            ).reshape((pix.height, pix.width, 3))

            ocr_text = extract_text_from_image(img_np)

            if ocr_text.strip():

                page_records = (
                    parse_text_records(
                        ocr_text
                    )
                )

                ocr_records.extend(
                    page_records
                )

        document.close()

        if ocr_records:

            return ocr_records

        raise ValueError(
            "No student data could be extracted from the PDF."
        )

    # =====================================================
    # PDF OPEN / CORRUPTION ERROR
    # =====================================================

    except ImportError as exc:

        raise ValueError(
            (
                "PDF support requires PyMuPDF. "
                "Run: pip install pymupdf"
            )
        ) from exc

    except Exception as exc:

        # -------------------------------------------------
        # Last fallback: try pypdf
        # -------------------------------------------------

        try:

            if PdfReader is not None:

                reader = PdfReader(
                    file_path,
                    strict=False,
                )

                fallback_pages = []

                for page in reader.pages:

                    try:

                        page_text = (
                            page.extract_text()
                            or ""
                        )

                        if page_text.strip():

                            fallback_pages.append(
                                page_text
                            )

                    except Exception:
                        continue

                fallback_text = "\n".join(
                    fallback_pages
                )

                if fallback_text.strip():

                    records = (
                        parse_text_records(
                            fallback_text
                        )
                    )

                    if records:
                        return records

        except Exception:
            pass

        raise ValueError(
            (
                "Unable to read PDF. "
                "The PDF may be corrupted, incomplete, "
                "or unsupported. "
                f"Details: {exc}"
            )
        ) from exc

        
# =========================================================
# MAIN FILE PARSER
# =========================================================

def parse_file(
    file_path: str,
):

    extension = (
        Path(file_path)
        .suffix
        .lower()
    )

    if extension == ".xlsx":

        return parse_xlsx(
            file_path
        )

    if extension == ".csv":

        return parse_csv(
            file_path
        )

    if extension == ".docx":

        return parse_docx(
            file_path
        )

    if extension == ".txt":

        return parse_txt(
            file_path
        )

    if extension in {
        ".jpg",
        ".jpeg",
        ".png",
        ".webp",
    }:

        return parse_image(
            file_path
        )

    if extension == ".pdf":

        return parse_pdf(
            file_path
        )

    raise ValueError(
        (
            "Unsupported file type. "
            "Supported: XLSX, CSV, DOCX, TXT, "
            "PDF, JPG, JPEG, PNG, WEBP"
        )
    )